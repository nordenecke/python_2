# -*- coding: utf-8 -*-
"""
Created on Sat Jul 15 21:58:21 2017

@author: norden
"""

import os
import word_content_structure
import traceback

from docx import Document
from docx.shared import Pt
#from docx.shared import Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
#from docx.oxml.shared import OxmlElement, qn

import _winreg
def get_desktop():
    key = _winreg.OpenKey(_winreg.HKEY_CURRENT_USER,\
                          r'Software\Microsoft\Windows\CurrentVersion\Explorer\Shell Folders',)
    return _winreg.QueryValueEx(key, "Desktop")[0]


words_list=[]
c_lst=[]

output_column_number=3
output_row_number=5



def get_word_list(filename):
    desktop_path=get_desktop()
    with open(desktop_path+"\\"+filename,"r") as words_src:
        for line in words_src:
            words_list.append(line.strip())
#            print(line.strip())
#        print(words_list)
        return words_list

def put_content(filename, content_list):
    with open(filename,"w+") as content_dst:
        content_dst.truncate()
        for i in range(len(content_list)):
            content_dst.write(content_list[i].word+"\n")
            content_dst.write(content_list[i].phonetic_symbol+"\n")
            content_dst.write(content_list[i].paraphrase+"\n")
            content_dst.write("***************************\n")
        return True
    
def add_word_and_phonetic_symbol_paragraph(cell,word,phonetic_symbol):
    #add word
    paragraph1=cell.add_paragraph(u"")
    paragraph1.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph1.add_run(word)
    if len(word) <= 7:
        run.font.size = Pt(48)
    elif len(word) <= 9:
        run.font.size = Pt(36)
    elif len(word) <= 12:
        run.font.size = Pt(28)
    elif len(word) <= 14:
        run.font.size = Pt(24)
    elif len(word) <= 15:
        run.font.size = Pt(22)
    elif len(word) <= 16:
        run.font.size = Pt(20)
    elif len(word) <= 18:
        run.font.size = Pt(18)
    else:
        run.font.size = Pt(16)
        
    run.font.name = 'Consolas'
    run.bold=True


    #add phonetic_symbol
    paragraph2=cell.add_paragraph(u"")
    paragraph2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph2.add_run(phonetic_symbol)
    run.font.size = Pt(16)
    run.font.name = 'Consolas'
    run.bold=True

    
def add_paraphrase_paragraph(cell,paraphrase):
    #add word
    paragraph1=cell.add_paragraph(u"")
    paragraph1.alignment=WD_ALIGN_PARAGRAPH.LEFT
    #set font size
    run = paragraph1.add_run(paraphrase)
    run.font.size = Pt(12)
    run.font.name=u'微软雅黑'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

    #set font
#    run = paragraph1.add_run(u'Set Font.')
#    run.font.name = 'Consolas'

    #set chinese font
#    run = paragraph1.add_run(u'set chinese font')
#    run.font.name=u'宋体'
#    r = run._element
#    r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    #set italic
#    run = paragraph1.add_run(u'set italic、')
#    run.italic = True

    #set bold
#    run = paragraph1.add_run(u'bold').bold = False
    
def set_raws_height(rows):
    for i in range(len(rows)):
        tr=rows[i]._tr
        trPr=tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), "3032")
        trHeight.set(qn('w:hRule'), "atLeast")
        trPr.append(trHeight)  



def set_cell_vertical_alignment(cell, align="center"): 
    try:   
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcValign = OxmlElement('w:vAlign')  
        tcValign.set(qn('w:val'), align)  
        tcPr.append(tcValign)
        return True 
    except:
        traceback.print_exc()             
        return False
    
def put_docx(filename,content_list):
    if len(content_list)==0:
        return False

    #open document
    document = Document()
    #加入不同等级的标题
#    document.add_heading(u'MS WORD写入测试',0)
#    document.add_heading(u'一级标题',1)
#    document.add_heading(u'二级标题',2)
    #添加文本
#    paragraph = document.add_paragraph(u'我们在做文本测试！')
    #设置字号
#    run = paragraph.add_run(u'设置字号、')
#    run.font.size = Pt(24)

    #设置字体
#    run = paragraph.add_run('Set Font,')
#    run.font.name = 'Consolas'

    #设置中文字体
#    run = paragraph.add_run(u'设置中文字体、')
#    run.font.name=u'宋体'
#    r = run._element
#    r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    #设置斜体
#    run = paragraph.add_run(u'斜体、')
#    run.italic = True

    #设置粗体
#    run = paragraph.add_run(u'粗体').bold = True

    #增加引用
#    document.add_paragraph('Intense quote', style='Intense Quote')

    #增加无序列表
#    document.add_paragraph(
#        u'无序列表元素1', style='List Bullet'
#    )
#    document.add_paragraph(
#        u'无序列表元素2', style='List Bullet'
#    )
    #增加有序列表
#    document.add_paragraph(
#        u'有序列表元素1', style='List Number'
#    )
#    document.add_paragraph(
#        u'有序列表元素2', style='List Number'
#    )
    #增加图像（此处用到图像image.bmp，请自行添加脚本所在目录中）
#    document.add_picture('image.bmp', width=Inches(1.25))

    page_item_number=output_column_number*output_row_number
    total_input_item_number=len(content_list)
    unused_item_number=total_input_item_number
    current_item=0
    cell_width=1828800*1.35

    while unused_item_number>0:
        #增加表格
        table1 = document.add_table(rows=output_row_number, cols=output_column_number, style='Table Grid')
        table1.autofit = False
        table1.alignment = WD_TABLE_ALIGNMENT.CENTER
        page_first_item=current_item
        for i in range(min(unused_item_number,page_item_number)):
#            first_page_content_item=content_list[current_item].word+"\n"+content_list[current_item].phonetic_symbol
            hdr_cells = table1.rows[int(i/output_column_number)].cells
#            hdr_cells[i%output_column_number].text = first_page_content_item
            add_word_and_phonetic_symbol_paragraph(hdr_cells[i%output_column_number],content_list[current_item].word,content_list[current_item].phonetic_symbol)
#            print hdr_cells[i%output_column_number].width
            set_cell_vertical_alignment(hdr_cells[i%output_column_number]) 
            hdr_cells[i%output_column_number].width=cell_width
            current_item+=1
        set_raws_height(table1.rows)
        #增加分页
        document.add_page_break()

        #增加表格
        table2 = document.add_table(rows=output_row_number, cols=output_column_number, style='Table Grid')
        table2.autofit = False
        table2.alignment = WD_TABLE_ALIGNMENT.CENTER
        current_item=page_first_item
        for i in range(min(unused_item_number,page_item_number)):
#            second_page_content_item=content_list[current_item].paraphrase
            hdr_cells = table2.rows[int(i/output_column_number)].cells
#            hdr_cells[output_column_number-i%output_column_number-1].text = second_page_content_item
            add_paraphrase_paragraph(hdr_cells[output_column_number-i%output_column_number-1],content_list[current_item].paraphrase)
            set_cell_vertical_alignment(hdr_cells[output_column_number-i%output_column_number-1]) 
            hdr_cells[output_column_number-i%output_column_number-1].width=cell_width
            current_item+=1
        set_raws_height(table2.rows)
        if unused_item_number -page_item_number<=0:
            break
        else:
            document.add_page_break()
            unused_item_number-=min(page_item_number,unused_item_number)

    #保存文件
    desktop_path=get_desktop()
    if os.path.exists(desktop_path+"\\"+filename):
        os.remove(desktop_path+"\\"+filename)
    Sections = document.sections
    for section in Sections:
        section.top_margin = 914400/8
        section.bottom_margin = 914400/8
        section.left_margin = 1143000/8
        section.right_margin = 1143000/8
    document.save(desktop_path+"\\"+filename)


def main():
#    print("This is get_word_list main function!")
#    print(get_word_list("words.txt"))
    c_lst.append(word_content_structure.word_content("111","bbb","ccc"))
    c_lst.append(word_content_structure.word_content("222","bbb","ccc"))
    put_content("words_content.txt",c_lst)


if __name__=="__main__":
    main()
    
    