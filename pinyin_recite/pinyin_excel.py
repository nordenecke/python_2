# -*- coding: utf-8 -*-
"""
Spyder Editor

This is a temporary script file.
"""

import os

import traceback

#import pypinyin
#from pypinyin import pinyin, lazy_pinyin
from pypinyin import pinyin

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
###############################################################################
input_word_list_file=u"Chinese_words.txt"
output_docx_file=u"看拼音写汉字.docx"
page_column_number = 20
Page_row_number = 16

###############################################################################

class AssertError(Exception):
    pass

class StopError(Exception):
    pass


def ASSERT(condition):
    if not condition:
        raise AssertError()

def STOP(message=''):
    raise StopError(message)

def load_cn_words(fname):
    ret=[]
    print 'Chinese file path='+fname
    if not os.path.isfile(fname):
        STOP('file %s not exist!'% fname)
    for l in open(fname,'r').read().decode(encoding='utf8').splitlines():
#        print l
#        l = l.split('#')[0].strip()
        if not l or l == '\ufeff':
            continue
        for w in l.split(' '):
            if not w or w.startswith('~'):
                continue
            if w in ret:
                continue
            ret.append(w)
    return ret

###############################################################################
import _winreg
def get_desktop():
    key = _winreg.OpenKey(_winreg.HKEY_CURRENT_USER,\
                          r'Software\Microsoft\Windows\CurrentVersion\Explorer\Shell Folders',)
    return _winreg.QueryValueEx(key, "Desktop")[0]


def get_cn_word_file_path(filename):
    desktop_path=get_desktop()
    file_path = desktop_path+"\\"+filename
    return file_path
###############################################################################

def set_cell_vertical_alignment(cell, align="center"):
    try:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcValign = OxmlElement('w:vAlign')
        tcValign.set(qn('w:val'), align)
        tcPr.append(tcValign)
        return True
    except:
        traceback.print_exc()
        return False

def generator(hzlst):
    if not hzlst:
        STOP()
    pinyin_lst = []
#    hanzi_lst = []
    source=list(hzlst)
    while True:
        if not source:
            break
        word = source.pop(0)
#        print '\nWord='+word
#        print pinyin(word)
#        pinyin_lst += [i[0] for i in pinyin(word)]
        if word.startswith('#'):
            pinyin_lst.append(word)
        else:
            pinyin_lst.append(pinyin(word))
#        hanzi_lst.append(word)
#        print 'word=['+word+']'
    # 选择输出内容
    #return pinyin_lst + ['EOL'] + hanzi_lst  # 拼音+汉字
    #return ['EOL'] + hanzi_lst  # 仅汉字
    #return pinyin_lst + ['EOL']  # 仅拼音
#    print '\npinyin_lst='
#    print pinyin_lst
#    print '\nhanzi_lst='
#    print hanzi_lst
    return pinyin_lst


def pinyin_generator():
    source = load_cn_words(get_cn_word_file_path(input_word_list_file))
    print '\nsource='
    print source

#    columns, rows = 15, 11  # 控制每页行列
#    page_limit = 20  # 控制输出页数
    return generator(source)

def add_lesson_number_paragraph(document,lesson_number):
    paragraph1 = document.add_paragraph(u'')

    #add word
    paragraph1.alignment=WD_ALIGN_PARAGRAPH.CENTER
    #set font size
    run = paragraph1.add_run(lesson_number)
    run.font.size = Pt(12)
    run.font.name=u'微软雅黑'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

def add_pinyin_paragraph(cell,lesson_number):
    paragraph1 = cell.add_paragraph(u'')

    #add word
    paragraph1.alignment=WD_ALIGN_PARAGRAPH.CENTER
    #set font size
    run = paragraph1.add_run(lesson_number)
    run.font.size = Pt(10)
    run.font.name=u'微软雅黑'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

def add_hanzi_space_paragraph(cell,hanzi_space_number):
    paragraph1 = cell.add_paragraph(u'')

    #add word
    paragraph1.alignment=WD_ALIGN_PARAGRAPH.CENTER
    #set font size
    hanzi_space=' '
    for i in range(hanzi_space_number):
        hanzi_space+=' '
    hanzi_space+=' '
    run = paragraph1.add_run(hanzi_space)
    run.font.size = Pt(10)
    run.font.name=u'微软雅黑'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

def create_twolines_table(document):
#    table = document.add_table(rows=2, cols=len(j))
    table = document.add_table(rows=2, cols=page_column_number, style='Table Grid')
#    table = document.add_table(rows=2, cols=len(j), style='Light Shading Accent 1')
    table.autofit = False
#    table.autofit = True
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    return table

def set_twolines_table_cell(table,lst):
    for i in range(len(lst)):
        py_cell=table.rows[0].cells
        hz_cell=table.rows[1].cells
        add_pinyin_paragraph(py_cell[i],lst[i])
        set_cell_vertical_alignment(py_cell[i])
        add_hanzi_space_paragraph(hz_cell[i],2)
    return

def get_py_one_word(lst, index):
    print 'lst[index]='+bytes(index)
    print lst[index]
    length=0
    for i in range(len(lst[index])):
      print len(lst[index][i][0])
      length+=len(lst[index][i][0])
      length+=1                     #add space
    if length > 0:
      length-=1
    print length
    return length

def get_hz_one_word(lst, index):
    print 'lst[index]='+bytes(index)
    print lst[index]
    length=0
    for i in range(len(lst[index])):
      print len(lst[index][i][0])
      length+=len(lst[index][i][0])*2
      length+=1                     #add space
    if length > 0:
      length-=1
    print length
    return length

def export2doc(hzlst, pylst, export_file_name):
    if len(hzlst)==0 or len(pylst)==0:
        return False

    first = True
    #open document
    document = Document()

    column_width=0
    begin = 0
    for j in range(len(pylst)):
        print j
        if pylst[j][0][0].startswith('#'):
            lesson_number= u'第'+hzlst[j][1:]+u'课'
            print lesson_number
            if first == True:
                print 'Do nothing for the first #!'
                first = False
                add_lesson_number_paragraph(document,lesson_number)
            else:
                document.add_page_break()
                add_lesson_number_paragraph(document,lesson_number)
        else:
            hzlen=get_hz_one_word(hzlst,j)
            pylen=get_py_one_word(pylst,j)
            length=max(hzlen,pylen)
            if column_width+length <= page_column_number:
                print length
            else:
                table=create_twolines_table(document)
                temp_list=pylst[begin:j-1]
                set_twolines_table_cell(table,temp_list)
                begin=j
#            add_lesson_number_paragraph(document,u'#########')

     #save file
    desktop_path=get_desktop()
    if os.path.exists(desktop_path+"\\"+export_file_name):
        os.remove(desktop_path+"\\"+export_file_name)
    Sections = document.sections
    for section in Sections:
        section.top_margin = 914400/8
        section.bottom_margin = 914400/8
        section.left_margin = 1143000/8
        section.right_margin = 1143000/8
    document.save(desktop_path+"\\"+export_file_name)
    return True

def main():
    hzlst = load_cn_words(get_cn_word_file_path(input_word_list_file))
    print '\nhzlst='
    print hzlst

#    columns, rows = 15, 11  # 控制每页行列
#    page_limit = 20  # 控制输出页数
    pylst=generator(hzlst)
    print '\npylst='
    print pylst
    export2doc(hzlst,pylst,output_docx_file)

if __name__=="__main__":
    main()
